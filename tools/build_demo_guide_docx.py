from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"D:\IOTs\projects\health-band-simulation")
OUTPUT = Path(r"D:\IOTs\tailieu26\HuongDan_Demo_DayDu_VongTaySucKhoe.docx")
EVIDENCE = PROJECT_ROOT / "tests" / "evidence"
CURRENT = EVIDENCE / "current-dashboard"

NAVY = "173B69"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1E7A4D"
GOLD = "7A5A00"
RED = "9B1C1C"
MUTED = "5D6B7B"
WHITE = "FFFFFF"
BLACK = "111827"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 80
CELL_SIDE = 120


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_SIDE, bottom=CELL_TOP_BOTTOM, end=CELL_SIDE):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths_dxa, header=True, font_size=9.2):
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    if header:
        mark_header_row(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_shading(cell, PALE_BLUE)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.12
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(header and row_index == 0), color=BLACK)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_picture_alt(paragraph, description):
    drawings = paragraph._p.xpath(".//wp:docPr")
    if drawings:
        drawings[-1].set("descr", description)
        drawings[-1].set("title", description)


def add_numbering_definition(doc, ordered):
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing_abstract or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_num or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_list_item(doc, text, num_id, bold_prefix=None):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True, color=BLACK)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=BLACK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = paragraph.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=BLACK)
        r2 = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=BLACK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_callout(doc, label, text, tone="blue"):
    fills = {"blue": CALLOUT, "green": "EDF8F2", "gold": "FFF8E8", "red": "FDEEEF"}
    colors = {"blue": DARK_BLUE, "green": GREEN, "gold": GOLD, "red": RED}
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [CONTENT_DXA], header=False, font_size=10)
    set_cell_shading(table.cell(0, 0), fills[tone])
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    r1 = paragraph.add_run(f"{label}: ")
    set_run_font(r1, size=10.2, bold=True, color=colors[tone])
    r2 = paragraph.add_run(text)
    set_run_font(r2, size=10.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, italic=True, color=MUTED)
    return paragraph


def add_figure(doc, path, caption, width=6.3):
    path = Path(path)
    if not path.exists():
        add_callout(doc, "Thiếu ảnh", f"Không tìm thấy {path.name}", "red")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_picture_alt(paragraph, caption)
    add_caption(doc, caption)


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)


def set_header_footer(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        r1 = hp.add_run("HEALTH BAND 01")
        set_run_font(r1, size=9, bold=True, color=BLUE)
        r2 = hp.add_run("  |  HƯỚNG DẪN DEMO")
        set_run_font(r2, size=9, color=MUTED)
        p_pr = hp._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "D7DEE8")
        borders.append(bottom)
        p_pr.append(borders)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run("Đề tài 31 · Nhóm Thành Danh  |  Trang ")
        set_run_font(r, size=9, color=MUTED)
        add_page_field(fp, "PAGE")
        r2 = fp.add_run("/")
        set_run_font(r2, size=9, color=MUTED)
        add_page_field(fp, "NUMPAGES")


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    widths = [2340, 2340, 2340, 2340]
    values = [
        ("12–15 phút", "Thời lượng đề xuất"),
        ("6 trang", "Dashboard hiện tại"),
        ("4 tầng", "Kiến trúc IoT"),
        ("13 tình huống", "Điểm demo chính"),
    ]
    for i, (value, label) in enumerate(values):
        cell = table.cell(0, i)
        set_cell_shading(cell, "EEF4FD")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(value)
        set_run_font(r1, size=13, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, color=MUTED)
    style_table(table, widths, header=False, font_size=9)


def add_page_demo(doc, title, objective, actions, explanation, evidence, image_path, figure_no):
    heading = add_heading(doc, title, 1)
    heading.paragraph_format.page_break_before = True
    add_callout(doc, "Mục tiêu", objective, "blue")
    add_heading(doc, "Thao tác trình bày", 2)
    for item in actions:
        add_list_item(doc, item, ORDERED_NUM_ID)
    add_heading(doc, "Cách chức năng hoạt động", 2)
    add_body(doc, explanation)
    add_callout(doc, "Bằng chứng phải chỉ ra", evidence, "green")
    add_figure(doc, image_path, f"Hình {figure_no}. {title}", width=6.25)


def build():
    global BULLET_NUM_ID, ORDERED_NUM_ID
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)
    set_header_footer(doc)
    BULLET_NUM_ID = add_numbering_definition(doc, ordered=False)
    ORDERED_NUM_ID = add_numbering_definition(doc, ordered=True)

    props = doc.core_properties
    props.title = "Hướng dẫn demo đầy đủ - Vòng tay theo dõi sức khỏe cá nhân"
    props.subject = "Kịch bản thuyết trình và giải thích hoạt động hệ thống IoT"
    props.author = "Nhóm Thành Danh"
    props.keywords = "IoT, ESP32, Wokwi, MQTT, Node-RED, Health Band"
    props.comments = "Tài liệu được tạo từ trạng thái dự án hiện hành."

    # Cover / workshop agenda header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("KỊCH BẢN THUYẾT TRÌNH & DEMO")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Vòng tay theo dõi\nsức khỏe cá nhân")
    set_run_font(r, size=29, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Đề tài 31 · ESP32/Wokwi · MQTT · Node-RED Dashboard 2.0")
    set_run_font(r, size=13.2, color=MUTED)

    add_metric_strip(doc)
    doc.add_paragraph()
    add_callout(
        doc,
        "Kết quả cần đạt",
        "Người xem thấy được dữ liệu đi từ thiết bị mô phỏng đến Dashboard, lệnh đi ngược về ESP32, hệ thống xử lý cảnh báo, tạo báo cáo và phục hồi trạng thái sau sự cố.",
        "green",
    )

    add_heading(doc, "Thành viên và vai trò khi demo", 2)
    roles = [
        ("Thành Danh", "Nhóm trưởng; mở đầu, kết luận, Tổng quan và luồng MQTT hai chiều."),
        ("Hồng Vỹ", "Wokwi, cảm biến, Vòng tay số và đầu ra OLED/LED/buzzer/rung."),
        ("Minh Thiện", "Trợ lý thông minh, hồ sơ sức khỏe, BMI và mục tiêu cá nhân."),
        ("Lê Hậu", "Nhật ký hoạt động, báo cáo, sao lưu/phục hồi và kiểm tra bằng chứng."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Thành viên"
    table.rows[0].cells[1].text = "Phần trình bày chính"
    for name, role in roles:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = role
    style_table(table, [2200, 7160])

    add_body(doc, f"Phiên bản tài liệu: {datetime.now().strftime('%d/%m/%Y')}. Dashboard: http://localhost:1880/dashboard/overview")
    add_callout(
        doc,
        "Giới hạn",
        "Đây là mô phỏng phục vụ học tập. Các giá trị HR, SpO₂, huyết áp và phân tích giấc ngủ không dùng để chẩn đoán y tế.",
        "gold",
    )

    heading = add_heading(doc, "1. Nội dung tài liệu", 1)
    heading.paragraph_format.page_break_before = True
    contents = [
        "Mục tiêu và tiêu chí demo thành công.",
        "Checklist chuẩn bị hệ thống trước giờ trình bày.",
        "Kiến trúc bốn tầng và luồng dữ liệu MQTT hai chiều.",
        "Kịch bản thời gian 12–15 phút.",
        "Hướng dẫn demo chi tiết sáu trang Dashboard.",
        "Ma trận tình huống Normal, cảnh báo, Fall/SOS, Sleep, Off wrist và Recovery.",
        "Cơ chế báo cáo, sao lưu, Gemini và xử lý khi không đeo.",
        "Câu hỏi phản biện, phương án dự phòng và checklist ảnh bằng chứng.",
    ]
    for item in contents:
        add_list_item(doc, item, ORDERED_NUM_ID)

    add_heading(doc, "2. Bảng viết tắt", 1)
    abbrev = [
        ("IoT", "Internet of Things", "Mạng thiết bị có khả năng thu thập, trao đổi và xử lý dữ liệu."),
        ("MQTT", "Message Queuing Telemetry Transport", "Giao thức publish/subscribe nhẹ dùng giữa ESP32, broker và Node-RED."),
        ("HR", "Heart Rate", "Nhịp tim, đơn vị BPM."),
        ("SpO₂", "Peripheral Oxygen Saturation", "Nồng độ oxy ngoại vi trong máu, đơn vị %."),
        ("BPM", "Beats Per Minute", "Số nhịp tim trong một phút."),
        ("NVS", "Non-Volatile Storage", "Bộ nhớ không mất dữ liệu của ESP32 dùng cho checkpoint."),
        ("BMI", "Body Mass Index", "Chỉ số khối cơ thể: cân nặng chia bình phương chiều cao."),
        ("GPS", "Global Positioning System", "Vị trí giả lập qua dữ liệu NMEA."),
        ("JSON", "JavaScript Object Notation", "Định dạng payload trao đổi trên MQTT."),
        ("LWT", "Last Will and Testament", "Bản tin MQTT báo offline khi thiết bị mất kết nối bất ngờ."),
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, text in enumerate(("Viết tắt", "Tên đầy đủ", "Ý nghĩa trong đồ án")):
        table.rows[0].cells[i].text = text
    for code, full, meaning in abbrev:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = code, full, meaning
    style_table(table, [1200, 3000, 5160], font_size=8.8)

    add_heading(doc, "3. Checklist trước giờ demo", 1)
    checklist = [
        "Docker Desktop đang chạy; container health-band-node-red ở trạng thái healthy.",
        "Mở http://localhost:1880/dashboard/overview và nhận HTTP 200.",
        "Wokwi Simulator đã Start; OLED/LED hiển thị và firmware không restart liên tục.",
        "Node-RED MQTT In/Out hiển thị connected với broker.emqx.io:1883.",
        "Dashboard hiển thị Đã kết nối và dữ liệu thay đổi khoảng 2 giây/lần ở Live mode.",
        "Chọn Normal trước khi bắt đầu để xóa tình huống bất thường.",
        "Chuẩn bị ảnh trong tests/evidence và video dự phòng nếu Internet hoặc broker lỗi.",
        "Không hiển thị file .env, API key Gemini hoặc dữ liệu cá nhân thật lên màn chiếu.",
    ]
    for item in checklist:
        add_list_item(doc, item, BULLET_NUM_ID)

    add_heading(doc, "3.1. Lệnh kiểm tra nhanh", 2)
    quick_checks = [
        ("docker compose ps", "Container Node-RED ở trạng thái Up (healthy)."),
        ("Test-NetConnection broker.emqx.io -Port 1883", "TcpTestSucceeded = True."),
        ("Mở /dashboard/overview", "Trang trả HTTP 200 và hiện trạng thái Connected."),
        ("Build firmware PlatformIO", "Kết thúc bằng [SUCCESS], tạo firmware.bin."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Kiểm tra"
    table.rows[0].cells[1].text = "Kết quả đạt"
    for command, expected in quick_checks:
        cells = table.add_row().cells
        cells[0].text = command
        cells[1].text = expected
    style_table(table, [4200, 5160], font_size=8.8)

    heading = add_heading(doc, "4. Hệ thống hoạt động như thế nào", 1)
    heading.paragraph_format.page_break_before = True
    add_body(
        doc,
        "Hệ thống vẫn dùng kiến trúc IoT bốn tầng dù trang Kiến trúc IoT đã được loại khỏi menu ứng dụng. Khi thuyết trình, giải thích kiến trúc bằng sơ đồ dưới đây và đối chiếu với Wokwi, Node-RED và Dashboard.",
    )
    add_figure(
        doc,
        Path(r"D:\IOTs\tailieu\do-an-31\docs\02-so-do-mo-hinh\kien-truc-iot-4-tang.png"),
        "Hình 1. Kiến trúc IoT bốn tầng của hệ thống Health Band",
        width=6.35,
    )
    architecture = [
        "Tầng 1 – Thiết bị: ESP32 đọc đầu vào HR/SpO₂, MPU6050, DS18B20, LDR, BMP180 và GPS giả lập; đồng thời điều khiển OLED, LED RGB, buzzer và phản hồi rung.",
        "Tầng 2 – Mạng: ESP32 kết nối Wokwi-GUEST và trao đổi dữ liệu với broker EMQX bằng MQTT.",
        "Tầng 3 – Xử lý: Node-RED parse JSON, kiểm tra chất lượng, phát hiện offline, xác nhận bất thường sau nhiều mẫu và phát cảnh báo.",
        "Tầng 4 – Ứng dụng: Dashboard hiển thị dữ liệu, gửi command, tạo báo cáo, lưu lịch sử và hỗ trợ sao lưu/phục hồi.",
    ]
    for item in architecture:
        add_list_item(doc, item, ORDERED_NUM_ID)

    add_heading(doc, "4.1. Luồng telemetry và command", 2)
    data_flow = [
        "ESP32 đọc hoặc sinh dữ liệu, đóng gói JSON và publish topic telemetry.",
        "Broker chuyển bản tin tới Node-RED mà không phụ thuộc giao diện web.",
        "Node-RED kiểm tra JSON, cập nhật thời điểm last seen, đánh giá rule rồi đẩy dữ liệu sang Dashboard.",
        "Khi người xem chọn một tình huống, Dashboard gửi command qua Node-RED và MQTT về ESP32.",
        "ESP32 đổi mode, cập nhật OLED/LED/buzzer/rung và publish event xác nhận; telemetry mới quay lại Dashboard.",
    ]
    for item in data_flow:
        add_list_item(doc, item, ORDERED_NUM_ID)
    add_callout(
        doc,
        "Điểm phải nhấn mạnh",
        "Dashboard không tự đổi số để giả vờ. Nút bấm gửi command thật; kết quả chỉ được xác nhận khi event/telemetry từ ESP32 quay lại.",
        "blue",
    )

    add_heading(doc, "4.2. Hợp đồng MQTT", 2)
    topics = [
        (".../telemetry", "ESP32 → Node-RED", "Dữ liệu định kỳ 2 giây Live hoặc 8 giây Eco."),
        (".../status", "ESP32/LWT → Node-RED", "Online/offline, retained, QoS 1."),
        (".../command", "Dashboard/Node-RED → ESP32", "Đổi mode, profile, power, reset, ack."),
        (".../event", "ESP32 → Node-RED", "Xác nhận command, FALL/SOS, phục hồi."),
        (".../alert", "Node-RED → Dashboard", "Kết quả rule và mức cảnh báo."),
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, text in enumerate(("Topic", "Hướng", "Vai trò")):
        table.rows[0].cells[i].text = text
    for topic, direction, role in topics:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = topic, direction, role
    style_table(table, [2200, 2600, 4560], font_size=9)
    add_body(doc, "Prefix chung: iot31/nhom-thanh-danh/health-band. Broker public chỉ dùng cho dữ liệu mô phỏng.")

    add_figure(
        doc,
        EVIDENCE / "TC01" / "TC01-nodered-mqtt-connected.png",
        "Hình 2. Luồng Node-RED nhận telemetry, đánh giá cảnh báo và gửi command",
        width=6.25,
    )

    heading = add_heading(doc, "5. Kịch bản thời gian 12–15 phút", 1)
    heading.paragraph_format.page_break_before = True
    schedule = [
        ("0:00–1:00", "Mở đầu", "Nêu bài toán, phạm vi mô phỏng và kiến trúc bốn tầng.", "Thành Danh"),
        ("1:00–3:00", "Thiết bị", "Chỉ ESP32, cảm biến, OLED, LED, buzzer, FALL/SOS.", "Hồng Vỹ"),
        ("3:00–5:00", "Tổng quan", "Normal, chỉ số chính, cảm biến mở rộng, biểu đồ.", "Thành Danh"),
        ("5:00–7:00", "Vòng tay số", "So sánh thiết bị ảo với Wokwi; demo Fall/SOS.", "Hồng Vỹ"),
        ("7:00–9:00", "Thông minh", "Điểm sức khỏe, mục tiêu, xu hướng, 3 mẫu, Eco.", "Minh Thiện"),
        ("9:00–10:30", "Hồ sơ", "BMI, profile trẻ em/người lớn, mục tiêu ngủ.", "Minh Thiện"),
        ("10:30–12:30", "Nhật ký", "Báo cáo, sao lưu/phục hồi, hoạt động gần đây.", "Lê Hậu"),
        ("12:30–15:00", "Kết luận/Q&A", "Recovery, giới hạn và câu hỏi phản biện.", "Cả nhóm"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(("Thời gian", "Phần", "Nội dung", "Người nói")):
        table.rows[0].cells[i].text = text
    for row in schedule:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [1350, 1500, 5010, 1500], font_size=8.8)
    add_callout(
        doc,
        "Nếu chỉ có 8–10 phút",
        "Bỏ phần nhập lại hồ sơ, chỉ giới thiệu nhanh App Guide; giữ bắt buộc Normal, High HR/Low SpO₂, Fall/SOS, report và recovery.",
        "gold",
    )

    add_page_demo(
        doc,
        "6. Trang Tổng quan",
        "Chứng minh thiết bị đang gửi dữ liệu, toàn bộ chỉ số được trình bày đúng thứ tự và người xem có thể điều khiển tình huống.",
        [
            "Chỉ trạng thái Connected và thời điểm đồng bộ gần nhất.",
            "Giới thiệu HR, SpO₂, huyết áp ước tính, bước chân, pin và điểm sức khỏe.",
            "Chỉ nhóm cảm biến mở rộng ngay dưới Chỉ số chính: nhiệt độ, ánh sáng, áp suất/độ cao, GPS và rung.",
            "Chỉ Biểu đồ trực tiếp và Chi tiết thiết bị.",
            "Cuộn xuống cuối trang để chọn Normal rồi High HR hoặc Low SpO₂; chờ đủ ba mẫu xác nhận.",
            "Chỉ Cảnh báo sức khỏe và bấm Acknowledge nếu có.",
        ],
        "Telemetry mới cập nhật state trung tâm. Dashboard tính trạng thái, vẽ 40 mẫu gần nhất, còn Node-RED tạo alert từ rule. Các nút ở cuối trang gửi command về ESP32 thay vì chỉ đổi giao diện.",
        "Chỉ số thay đổi; command được xác nhận; cảnh báo xuất hiện đúng ngưỡng; sensor cards nằm trên biểu đồ; điều khiển/cảnh báo ở cuối trang.",
        CURRENT / "overview-vi.png",
        3,
    )

    add_page_demo(
        doc,
        "7. Trang Vòng tay số",
        "Tạo liên hệ trực quan giữa ESP32/Wokwi và chiếc vòng tay ảo trên Dashboard.",
        [
            "Đặt Wokwi và Vòng tay số cạnh nhau.",
            "So sánh HR, SpO₂, bước, huyết áp, pin và trạng thái ngủ trên màn hình.",
            "Chọn Trigger FALL/SOS hoặc bấm nút đỏ Wokwi.",
            "Chỉ LED đỏ, buzzer/rung, banner khẩn cấp và chuỗi Nguyên nhân–kết quả.",
            "Cuộn xuống để chọn đối tượng, giới tính, trạng thái đeo và mục tiêu ngủ.",
        ],
        "Vòng tay số đọc cùng telemetry với Dashboard. LED ảo lấy trạng thái mode/cảnh báo; danh sách cause-and-effect chỉ hoàn tất khi command, acknowledgement, telemetry và alert đã xuất hiện.",
        "Giá trị Wokwi và vòng tay số khớp; FALL tạo đầu ra đỏ/còi/rung; event xác nhận xuất hiện.",
        CURRENT / "twin-vi.png",
        4,
    )

    add_page_demo(
        doc,
        "8. Trang Trợ lý thông minh",
        "Chứng minh ứng dụng không chỉ hiển thị số mà còn phân tích có quy tắc, cá nhân hóa và kiểm tra chất lượng dữ liệu.",
        [
            "Giải thích điểm sức khỏe 0–100 theo Vitals, Safety, Activity và Device.",
            "Nhập mục tiêu bước chân cá nhân rồi Save goal.",
            "Chọn Student, Older adult, Athlete hoặc Child và giới tính.",
            "Chỉ xu hướng 10 mẫu, bộ đếm xác nhận bất thường 1/3 → 3/3.",
            "Chỉ Data quality, Smart suggestions, Live/Eco mode và Normal vs Current.",
            "Bấm Start automatic demo nếu muốn chạy chuỗi tình huống tự động.",
        ],
        "Các kết quả được tính bằng rule minh bạch trong Dashboard/Node-RED. Gemini là bổ sung tùy chọn: chỉ nhận dữ liệu tổng hợp đã ẩn danh và luôn có local fallback.",
        "Mục tiêu lưu sau reload; profile đổi ngưỡng; xu hướng và chất lượng dữ liệu có giải thích; auto demo gửi command thật.",
        CURRENT / "smart-vi.png",
        5,
    )

    add_page_demo(
        doc,
        "9. Trang Hồ sơ sức khỏe",
        "Cho phép người dùng nhập thông tin cơ bản và nhận phân tích sàng lọc minh bạch.",
        [
            "Nhập tuổi, chiều cao, cân nặng, mức vận động và giới tính.",
            "Bấm Lưu và phân tích.",
            "Chỉ công thức BMI và nhãn kết quả đối với người từ 20 tuổi.",
            "Đổi tuổi thành 12 để chứng minh ứng dụng không áp nhãn BMI người lớn cho trẻ em.",
            "Bấm Áp dụng làm mục tiêu ngủ để cập nhật cấu hình cá nhân.",
        ],
        "BMI = cân nặng (kg) / chiều cao² (m). Với 5–19 tuổi, ứng dụng chỉ hiển thị giá trị và yêu cầu phân vị theo tuổi/giới tính. Dữ liệu hồ sơ lưu trong localStorage và bản backup, không publish lên broker public.",
        "BMI tính đúng; trẻ em hiển thị giới hạn phân vị; mục tiêu ngủ thay đổi; thông tin tồn tại sau reload.",
        CURRENT / "profile-vi.png",
        6,
    )

    add_page_demo(
        doc,
        "10. Trang Hướng dẫn sử dụng",
        "Giúp thành viên khác hoặc người xem hiểu từng trang và cách kiểm chứng chức năng ngay trong ứng dụng.",
        [
            "Chỉ các thẻ Tổng quan, Vòng tay số, Trợ lý thông minh, Hồ sơ sức khỏe và Nhật ký hoạt động.",
            "Đọc phần Mục đích, Tính năng và Cách trình diễn của một thẻ.",
            "Bấm Mở trang này để chứng minh điều hướng trực tiếp.",
            "Chỉ bảng giải thích chức năng ở cuối trang.",
        ],
        "Guide pages được tạo từ cấu hình trong Dashboard, vì vậy mỗi thẻ có target mở đúng trang. Đây là tài liệu hỗ trợ demo, không thay thế README hoặc báo cáo kỹ thuật.",
        "Không còn thẻ Kiến trúc IoT; có thẻ Nhật ký hoạt động; nút mở trang điều hướng đúng.",
        CURRENT / "guide-vi.png",
        7,
    )

    add_page_demo(
        doc,
        "11. Trang Nhật ký hoạt động",
        "Tập trung toàn bộ bằng chứng lịch sử, báo cáo và khả năng phục hồi tại một trang riêng.",
        [
            "Bấm Tạo ngay trong Báo cáo tự động; chỉ trạng thái theo giờ/ngày/tháng.",
            "Tải báo cáo JSON và chỉ wearCoveragePercent, số mẫu hợp lệ và giá trị trung bình.",
            "Bấm Xuất bản sao lưu để tải snapshot Dashboard.",
            "Giải thích ba checkpoint ESP32 NVS, Browser và Docker.",
            "Chỉ Hoạt động gần đây: command, acknowledgement, network, telemetry và alert.",
        ],
        "Báo cáo chỉ tổng hợp mẫu hợp lệ khi đang đeo; thời gian tháo vòng được ghi là thiếu coverage. Backup chứa cấu hình/trạng thái Dashboard; ESP32 và Docker có cơ chế lưu độc lập.",
        "File report/backup tải được; checkpoint có thời gian; timeline giữ sự kiện sau khi trở lại Normal.",
        CURRENT / "activity-vi.png",
        8,
    )

    heading = add_heading(doc, "12. Ma trận tình huống phải demo", 1)
    heading.paragraph_format.page_break_before = True
    scenarios = [
        ("Normal", "Chọn Normal", "HR 70–90; SpO₂ 97–99; LED xanh", "Không có alert; baseline được ghi"),
        ("High HR", "Chọn High heart rate, chờ 3 mẫu", "HR 125–145; LED cảnh báo", "Bộ đếm 1/3→3/3; alert HR cao"),
        ("Low HR", "Chọn Low heart rate, chờ 3 mẫu", "HR 40–48", "Alert HR thấp sau xác nhận"),
        ("Low SpO₂", "Chọn Low SpO₂, chờ 3 mẫu", "SpO₂ 85–92", "Alert oxy thấp"),
        ("Fall/SOS", "Chọn Fall hoặc nút đỏ Wokwi", "LED đỏ; buzzer/rung; FALL", "Banner đếm ngược; alert critical; GPS"),
        ("Low battery", "Chọn Low battery", "Pin 12–19%; Eco", "Cảnh báo pin; chu kỳ 8 giây"),
        ("Sleep", "Chọn Sleep, đợi 15–60 giây", "Ít chuyển động; trạng thái Zz", "Candidate → Light → Deep"),
        ("Off wrist", "Chọn Tháo khỏi tay", "Vitals hiển thị --/0", "Connected; không tính mẫu/report/sleep"),
        ("Reset steps", "Bấm Reset step counter", "Steps = 0", "Có event acknowledgement"),
        ("Offline", "Stop Wokwi >8 giây", "Thiết bị ngừng gửi", "Banner offline; LWT/status/timer"),
        ("Recovery", "Start Wokwi lại", "Khôi phục NVS", "DEVICE_RECOVERED; timeline giữ lịch sử"),
        ("Language/theme", "Đổi VN/US và sáng/tối", "Không đổi firmware", "UI đổi ngay; lưu localStorage"),
        ("Gemini", "Ask Gemini", "Không tác động thiết bị", "Khuyến cáo 3 dòng hoặc fallback"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(("Tình huống", "Thao tác", "Kết quả Wokwi", "Kết quả Dashboard")):
        table.rows[0].cells[i].text = text
    for scenario in scenarios:
        cells = table.add_row().cells
        for i, value in enumerate(scenario):
            cells[i].text = value
    style_table(table, [1500, 2400, 2600, 2860], font_size=8.2)

    add_heading(doc, "12.1. Demo Fall/SOS an toàn", 2)
    for item in [
        "Kích hoạt Fall và nói rõ đây là emergency simulation.",
        "Chỉ GPS giả lập, countdown 10 giây, LED đỏ, buzzer/rung và alert critical.",
        "Bấm Tôi an toàn để hủy hoặc chờ trạng thái simulated notification sent.",
        "Không nói rằng hệ thống đã gọi dịch vụ cấp cứu thật.",
    ]:
        add_list_item(doc, item, ORDERED_NUM_ID)

    add_heading(doc, "12.2. Demo theo dõi giấc ngủ", 2)
    sleep_steps = [
        "Kiểm tra đang đeo: dữ liệu thô HR/SpO₂ và nhiệt độ bề mặt phải phù hợp.",
        "Kiểm tra khung 21:00–09:00; mode Sleep cho phép demo ngoài khung.",
        "Sensor fusion yêu cầu không chuyển động, HR 45–70 BPM, SpO₂ ≥90% và nhiệt độ ≥30°C.",
        "Dưới 15 giây: candidate 45%; 15–30 giây: candidate 70%; từ 30 giây: light 85%; từ 60 giây: deep 92%.",
        "Nếu chuyển động hoặc tháo vòng, thuật toán reset về awake/not_tracked.",
    ]
    for item in sleep_steps:
        add_list_item(doc, item, ORDERED_NUM_ID)
    add_callout(
        doc,
        "Giới hạn kỹ thuật",
        "Ngưỡng 15–60 giây được tăng tốc cho lớp học. Thiết bị thật cần quan sát dài hơn, HRV/PPG ổn định và thuật toán được kiểm định; dự án không nhận diện REM.",
        "gold",
    )

    add_heading(doc, "13. Báo cáo, sao lưu và phục hồi", 1)
    add_heading(doc, "13.1. Báo cáo tự động", 2)
    reports = [
        "Theo giờ: tạo khi đủ một giờ dữ liệu đeo hợp lệ hoặc khi người dùng bấm Tạo ngay.",
        "Theo ngày: tổng hợp hoạt động trong ngày; thời gian không đeo không biến thành chỉ số sức khỏe bằng 0.",
        "Theo tháng: tổng hợp các phiên/ngày đã có dữ liệu.",
        "Trường quan trọng khi trình bày: sampleCount, validSampleCount, wearCoveragePercent, averages và alerts.",
    ]
    for item in reports:
        add_list_item(doc, item, BULLET_NUM_ID)

    add_heading(doc, "13.2. Ba lớp checkpoint", 2)
    checkpoints = [
        ("ESP32 NVS", "Lưu bước, pin, profile, giới tính và mode khoảng 30 giây/lần và sau command. Khi khởi động lại, firmware đọc checkpoint rồi phát event phục hồi."),
        ("Browser", "Lưu snapshot hiện tại và snapshot trước trong localStorage; chứa hồ sơ, mục tiêu, báo cáo, timeline và cấu hình Dashboard."),
        ("Docker/Node-RED", "Dữ liệu flow/context nằm trong volume gắn với container, nên recreate container không làm mất cấu hình."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Lớp"
    table.rows[0].cells[1].text = "Cách hoạt động"
    for name, detail in checkpoints:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = name, detail
    style_table(table, [1900, 7460], font_size=9)

    add_heading(doc, "13.3. Khi thiết bị sập nguồn", 2)
    for item in [
        "MQTT LWT hoặc bộ đếm 8 giây làm Dashboard chuyển Offline.",
        "Khi ESP32 khởi động lại, Preferences/NVS phục hồi checkpoint gần nhất.",
        "Thiết bị kết nối Wi‑Fi/MQTT, publish status online retained và event DEVICE_RECOVERED.",
        "Telemetry tiếp tục; Dashboard chuyển Online nhưng timeline vẫn giữ dấu vết sự cố.",
        "Dữ liệu giữa checkpoint cuối và lúc mất nguồn có thể mất; cần nói rõ giới hạn này.",
    ]:
        add_list_item(doc, item, ORDERED_NUM_ID)

    add_heading(doc, "14. Khi người dùng không đeo 24/24", 1)
    off_wrist = [
        "Thiết bị vẫn online để giữ command, bước chân, pin và trạng thái kết nối.",
        "HR, SpO₂, huyết áp và fall không được xem là mẫu hợp lệ khi off wrist.",
        "Giấc ngủ chuyển not_tracked; khoảng thời gian đó ghi là thiếu coverage.",
        "Báo cáo không thay dữ liệu thiếu bằng 0 vì điều đó làm sai trung bình.",
        "Khi đeo lại, cảm biến đủ điều kiện thì hệ thống tiếp tục thu mẫu; không tự suy diễn dữ liệu trong khoảng trống.",
    ]
    for item in off_wrist:
        add_list_item(doc, item, BULLET_NUM_ID)

    add_heading(doc, "15. Gemini và dữ liệu cá nhân", 1)
    add_body(
        doc,
        "Gemini được gọi qua endpoint Node-RED để API key không xuất hiện trong trình duyệt. Request chỉ gửi dữ liệu tổng hợp ẩn danh; không gửi tên, GPS, device ID hoặc lịch sử thô. Nếu API lỗi hoặc thiếu key, Dashboard dùng khuyến cáo cục bộ.",
    )
    add_callout(
        doc,
        "Khi demo",
        "Chỉ trình bày Gemini sau khi các rule cục bộ đã chạy. Nhấn mạnh AI chỉ gợi ý giáo dục, không chẩn đoán và không thay thế logic cảnh báo cốt lõi.",
        "blue",
    )

    heading = add_heading(doc, "16. Câu hỏi phản biện và trả lời ngắn", 1)
    heading.paragraph_format.page_break_before = True
    qa = [
        ("Vì sao chọn kiến trúc 4 tầng?", "Đủ tách thiết bị, mạng, xử lý và ứng dụng; dễ giải thích và phù hợp quy mô môn học."),
        ("Vì sao dùng MQTT?", "Nhẹ, publish/subscribe, tách rời ESP32 với Dashboard và hỗ trợ LWT/QoS/retain."),
        ("Dữ liệu có phải y tế thật không?", "Không. Cảm biến và kịch bản trên Wokwi chỉ dùng để chứng minh luồng IoT."),
        ("Tại sao chờ 3 mẫu mới cảnh báo?", "Giảm cảnh báo giả do một mẫu nhiễu; Fall vẫn được ưu tiên ngay."),
        ("Nếu tháo vòng thì sao?", "Thiết bị vẫn online nhưng mẫu sinh tồn/sleep bị loại, coverage được ghi lại."),
        ("Sập nguồn có mất hết không?", "Không; NVS, browser snapshot và Docker volume phục hồi phần lớn trạng thái. Có thể mất dữ liệu sau checkpoint cuối."),
        ("Gemini có quyết định cảnh báo không?", "Không. Rule cục bộ quyết định; Gemini chỉ tạo khuyến cáo tổng hợp tùy chọn."),
        ("Vì sao không còn trang Kiến trúc IoT?", "Giao diện ưu tiên người dùng; kiến trúc vẫn là nền tảng kỹ thuật và được trình bày trong báo cáo/slide."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Câu hỏi"
    table.rows[0].cells[1].text = "Trả lời đề xuất"
    for question, answer in qa:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = question, answer
    style_table(table, [3300, 6060], font_size=9)

    add_heading(doc, "17. Checklist ảnh và bằng chứng phải lưu", 1)
    evidence = [
        "Wokwi đang chạy: ESP32, OLED, LED RGB, buzzer/rung và FALL/SOS.",
        "Node-RED MQTT connected và flow parse/evaluate/publish.",
        "Tổng quan Normal có dữ liệu; cảm biến mở rộng nằm trên biểu đồ.",
        "High HR hoặc Low SpO₂ với bộ đếm xác nhận và alert.",
        "Fall/SOS có LED đỏ, countdown, GPS và critical alert.",
        "Vòng tay số khớp giá trị với Wokwi.",
        "Trợ lý thông minh có mục tiêu tùy chỉnh, xu hướng, data quality và auto demo.",
        "Hồ sơ sức khỏe có BMI người lớn và cảnh báo phân vị trẻ em.",
        "App Guide có đúng năm thẻ chức năng ngoài chính nó, không còn trang Kiến trúc.",
        "Nhật ký hoạt động có báo cáo, backup và timeline.",
        "Offline >8 giây và Recovery sau khi Start lại Wokwi.",
        "Giao diện tiếng Việt/English, theme sáng/tối và chuông thông báo.",
    ]
    for item in evidence:
        add_list_item(doc, f"☐ {item}", BULLET_NUM_ID)

    add_heading(doc, "18. Phương án dự phòng khi demo lỗi", 1)
    fallback = [
        ("Dashboard không có dữ liệu", "Ctrl+F5; kiểm tra Wokwi, MQTT connected và Internet; chọn Normal."),
        ("Broker public gián đoạn", "Dùng ảnh/video dự phòng; giải thích broker là thành phần bên ngoài."),
        ("Wokwi Terminal trống", "Dùng OLED, LED, Node-RED connected và Dashboard thay đổi làm bằng chứng."),
        ("Gemini lỗi", "Chỉ local fallback; không dành thời gian sửa API ngay trên lớp."),
        ("Container lỗi", "docker compose up -d --force-recreate; kiểm tra health và port 1880."),
        ("Thiếu thời gian", "Demo Normal → High HR/Low SpO₂ → Fall → Report → Recovery."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Sự cố"
    table.rows[0].cells[1].text = "Xử lý nhanh"
    for issue, resolution in fallback:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = issue, resolution
    style_table(table, [2700, 6660], font_size=9)

    add_heading(doc, "19. Kết luận mẫu", 1)
    add_callout(
        doc,
        "Lời kết 30 giây",
        "Dự án chứng minh một chu trình IoT hoàn chỉnh: thiết bị mô phỏng thu dữ liệu, MQTT vận chuyển hai chiều, Node-RED kiểm tra và tạo cảnh báo, còn Dashboard trực quan hóa, cá nhân hóa, báo cáo và phục hồi. Hệ thống ưu tiên khả năng giải thích và kiểm chứng; mọi giới hạn y tế đều được công bố rõ.",
        "green",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
